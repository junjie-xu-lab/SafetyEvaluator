"""Tests for SafetyEvaluator report generation."""

from io import BytesIO

import pandas as pd
import pytest

from safetyevaluator.metrics import calculate_detector_comparison
from safetyevaluator.report import (
    build_misclassified_samples_table,
    generate_multi_detector_excel_report,
    generate_multi_detector_html_report,
    generate_multi_detector_pdf_report,
    generate_multi_detector_word_report,
)


def test_excel_report_contains_expected_sheets() -> None:
    pytest.importorskip("openpyxl")

    data = pd.DataFrame(
        [
            {
                "id": "1",
                "input": "safe",
                "label": "Safe",
                "prediction_baseline": "Unsafe",
                "prediction_strict": "Safe",
                "source": "demo",
                "category": "normal",
            },
            {
                "id": "2",
                "input": "unsafe",
                "label": "Unsafe",
                "prediction_baseline": "Unsafe",
                "prediction_strict": "Unsafe",
                "source": "demo",
                "category": "cyber",
            },
        ]
    )
    comparison = calculate_detector_comparison(data, ["prediction_baseline", "prediction_strict"])

    report_bytes = generate_multi_detector_excel_report(comparison)
    workbook = pd.ExcelFile(BytesIO(report_bytes))

    assert workbook.sheet_names == [
        "Summary",
        "Detector Comparison",
        "Metrics",
        "Confusion Matrix",
        "Group Analysis",
        "Misclassified Samples",
    ]


def test_combined_misclassified_samples_include_detector_context() -> None:
    data = pd.DataFrame(
        [
            {"id": "1", "input": "safe", "label": "Safe", "prediction_a": "Unsafe", "prediction_b": "Safe"},
            {"id": "2", "input": "unsafe", "label": "Unsafe", "prediction_a": "Unsafe", "prediction_b": "Safe"},
        ]
    )
    comparison = calculate_detector_comparison(data, ["prediction_a", "prediction_b"])

    errors = build_misclassified_samples_table(comparison)

    assert list(errors["Prediction Column"]) == ["prediction_a", "prediction_b"]
    assert set(errors["error_type"]) == {"False Positive", "False Negative"}


def test_html_report_contains_detector_sections_and_escapes_values() -> None:
    data = pd.DataFrame(
        [
            {
                "id": "1",
                "input": "<script>alert('x')</script>",
                "label": "Safe",
                "prediction_a": "Unsafe",
                "source": "demo",
                "category": "normal",
            }
        ]
    )
    comparison = calculate_detector_comparison(data, ["prediction_a"])

    report_html = generate_multi_detector_html_report(comparison)

    assert report_html.startswith("<!DOCTYPE html>")
    assert "Detector: A" in report_html
    assert "&lt;script&gt;alert(&#x27;x&#x27;)&lt;/script&gt;" in report_html
    assert "<script>alert('x')</script>" not in report_html


def test_pdf_report_generates_pdf_bytes() -> None:
    pytest.importorskip("reportlab")

    data = pd.DataFrame(
        [
            {"id": "1", "input": "safe", "label": "Safe", "prediction": "Unsafe"},
            {"id": "2", "input": "unsafe", "label": "Unsafe", "prediction": "Unsafe"},
        ]
    )
    comparison = calculate_detector_comparison(data, ["prediction"])

    report_bytes = generate_multi_detector_pdf_report(comparison)

    assert report_bytes.startswith(b"%PDF")


def test_word_report_contains_core_sections() -> None:
    pytest.importorskip("docx")
    from docx import Document

    data = pd.DataFrame(
        [
            {
                "id": "1",
                "input": "safe",
                "label": "Safe",
                "prediction_baseline": "Unsafe",
                "prediction_strict": "Safe",
                "source": "demo",
                "category": "normal",
            },
            {
                "id": "2",
                "input": "unsafe",
                "label": "Unsafe",
                "prediction_baseline": "Unsafe",
                "prediction_strict": "Unsafe",
                "source": "demo",
                "category": "cyber",
            },
        ]
    )
    comparison = calculate_detector_comparison(data, ["prediction_baseline", "prediction_strict"])

    report_bytes = generate_multi_detector_word_report(comparison)
    document = Document(BytesIO(report_bytes))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert report_bytes.startswith(b"PK")
    assert "Safety Evaluation Report" in paragraph_text
    assert "Detector: Baseline" in paragraph_text
    assert "Detector: Strict" in paragraph_text
    assert any("Detector" in cell.text for table in document.tables for cell in table.rows[0].cells)
